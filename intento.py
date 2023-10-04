import os
from docx import Document
import subprocess

# Directorio base
base_directory = "/var/ossec/logs/alerts"

# Comando para listar archivos en el directorio con sudo
ls_command = f"sudo ls {base_directory}"

# Crear un archivo Word
doc = Document()
doc.add_heading("Alertas OSSEC", level=1)

# Obtener la lista de carpetas y archivos en el directorio con sudo
result = subprocess.run(ls_command, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

if result.returncode == 0:
    # Lista de archivos y carpetas en el directorio
    items = result.stdout.splitlines()
    # Recorrer las carpetas por año y mes
    for item in items:
        year_directory = os.path.join(base_directory, item)
        if os.path.isdir(year_directory):
            for month in os.listdir(year_directory):
                month_directory = os.path.join(year_directory, month)
                if os.path.isdir(month_directory):
                    # Buscar archivos que coincidan con el patrón "ossec-alerts-*.log" en esta carpeta
                    for filename in os.listdir(month_directory):
                        if filename.startswith("ossec-alerts-") and filename.endswith(".log"):
                            log_file = os.path.join(month_directory, filename)
                            with open(log_file, "r") as file:
                                # Leer el contenido del archivo de registro y agregarlo al archivo Word
                                doc.add_heading(f"{item}/{month}/{filename}", level=2)
                                doc.add_paragraph(file.read())

    # Guardar el archivo Word
    output_file = "alertas_ossec.docx"
    doc.save(output_file)

    print(f"Se han recopilado las alertas en {output_file}")
else:
    print(f"No se pudo obtener la lista de archivos en {base_directory}. Asegúrate de que tienes permisos de administrador y que has configurado sudo para permitir el acceso sin contraseña.")
